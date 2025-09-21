import os, json
from pathlib import Path
from typing import Dict, Any
import pandas as pd
from docx import Document

CONFIG = {
    "crescimento_receitas": [0.0, 0.08, 0.08],
    "margem_bruta_target": 0.55,
    "fse_pct_receitas": 0.12,
    "encargos_sociais_pct": 0.2375,
    "depreciacao_anos": {"equipamento":5,"informatica":3,"veiculos":4,"intangiveis":3,"outros":4},
    "taxa_juros": 0.06,
    "amortizacao_anos": 3,
}

def ia_disponivel():
    return bool(os.getenv("OPENAI_API_KEY"))

def gerar_texto_ia(titulo, instrucoes, contexto):
    if ia_disponivel():
        try:
            from openai import OpenAI
            client = OpenAI()
            prompt = f"Escreve um texto claro para '{titulo}' em PT-PT. Contexto: {json.dumps(contexto, ensure_ascii=False)}. {instrucoes}. 120-200 palavras."
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role":"user","content":prompt}],
                temperature=0.5,
            )
            return resp.choices[0].message.content.strip()
        except Exception:
            return f"[Gerar com IA] {titulo}: {instrucoes}"
    return f"[Preencher] {titulo}: {instrucoes}"

def aplicar_limite(texto: str, limite: int) -> str:
    if not texto or limite <= 0:
        return texto or ""
    if len(texto) <= limite:
        return texto
    recorte = texto[:limite]
    if " " in recorte:
        recorte = recorte.rsplit(" ", 1)[0]
    return recorte + "…"

def calcular_tabelas(dados: Dict[str, Any]) -> Dict[str, pd.DataFrame]:
    anos = dados.get("anos",[2025,2026,2027])
    y1,y2,y3 = anos
    # vendas
    vendas = []
    for v in dados.get("vendas",[]):
        y1t = v.get("preco",0)*v.get("qtd_mensal",0)*v.get("meses_y1",12)
        y2t = y1t*(1+CONFIG["crescimento_receitas"][1])
        y3t = y2t*(1+CONFIG["crescimento_receitas"][2])
        vendas.append({"designacao":v.get("designacao","—"), str(y1):round(y1t,2), str(y2):round(y2t,2), str(y3):round(y3t,2)})
    df_vendas = pd.DataFrame(vendas) if vendas else pd.DataFrame(columns=["designacao",str(y1),str(y2),str(y3)])
    tot = {str(y1): df_vendas[str(y1)].sum() if not df_vendas.empty else 0.0,
           str(y2): df_vendas[str(y2)].sum() if not df_vendas.empty else 0.0,
           str(y3): df_vendas[str(y3)].sum() if not df_vendas.empty else 0.0}

    cogs = {k: (1-CONFIG["margem_bruta_target"])*v for k,v in tot.items()}
    df_cogs = pd.DataFrame([{"rubrica":"COGS", **{k:round(v,2) for k,v in cogs.items()}}])
    fse  = {k: CONFIG["fse_pct_receitas"]*v for k,v in tot.items()}
    df_fse = pd.DataFrame([{"rubrica":"FSE", **{k:round(v,2) for k,v in fse.items()}}])

    # pessoal
    df_pessoal = pd.DataFrame(columns=["rubrica",str(y1),str(y2),str(y3)])
    if dados.get("pessoal"):
        linhas = []
        for p in dados["pessoal"]:
            meses = p.get("meses",12)
            y1b = p.get("venc_mensal",0)*p.get("n",0)*meses
            y2b = y1b*(1+dados.get("aumento_salarios_pct",0.03))
            y3b = y2b*(1+dados.get("aumento_salarios_pct",0.03))
            y1t = y1b*(1+CONFIG["encargos_sociais_pct"])
            y2t = y2b*(1+CONFIG["encargos_sociais_pct"])
            y3t = y3b*(1+CONFIG["encargos_sociais_pct"])
            linhas.append({"rubrica":p.get("funcao","—"), str(y1):round(y1t,2), str(y2):round(y2t,2), str(y3):round(y3t,2)})
        df_pessoal = pd.DataFrame(linhas)
    tot_pess = {str(y1): df_pessoal[str(y1)].sum() if not df_pessoal.empty else 0.0,
                str(y2): df_pessoal[str(y2)].sum() if not df_pessoal.empty else 0.0,
                str(y3): df_pessoal[str(y3)].sum() if not df_pessoal.empty else 0.0}

    # depreciacoes
    df_dep = pd.DataFrame(columns=["bem",str(y1),str(y2),str(y3)])
    dep_tot = {str(y1):0.0,str(y2):0.0,str(y3):0.0}
    for it in dados.get("investimento",[]):
        tipo = it.get("tipo","outros").lower()
        anos_dep = CONFIG["depreciacao_anos"].get(tipo,4)
        anu = float(it.get("valor",0))/anos_dep if anos_dep else 0.0
        df_dep.loc[len(df_dep)] = { "bem": f"{tipo}: {it.get('descricao','—')}", str(y1): round(anu,2), str(y2): round(anu,2), str(y3): round(anu,2)}
        for k in dep_tot: dep_tot[k]+=anu

    # financiamento
    df_fin = pd.DataFrame(columns=["ano","prestacao","capital","juros","divida_final"])
    juros_tot = {str(y1):0.0,str(y2):0.0,str(y3):0.0}
    emp = dados.get("emprestimo",{})
    if emp.get("montante",0)>0:
        principal=float(emp.get("montante",0)); anos_amort=int(emp.get("amortizacao_anos",CONFIG["amortizacao_anos"])); taxa=float(emp.get("taxa_juros",CONFIG["taxa_juros"]))
        amort=principal/anos_amort if anos_amort else principal
        saldo=principal
        linhas=[]
        for ano in [y1,y2,y3]:
            juros=saldo*taxa; capital=min(amort,saldo); prest=juros+capital; saldo=max(0.0,saldo-capital)
            juros_tot[str(ano)]=juros; linhas.append({"ano":ano,"prestacao":round(prest,2),"capital":round(capital,2),"juros":round(juros,2),"divida_final":round(saldo,2)})
        df_fin=pd.DataFrame(linhas)

    # DR
    df_dr = pd.DataFrame([
        {"rubrica":"Vendas/Serviços", str(y1):round(tot[str(y1)],2), str(y2):round(tot[str(y2)],2), str(y3):round(tot[str[y3]] if False else tot[str(y3)],2)},
        {"rubrica":"COGS", str(y1):round(df_cogs[str(y1)].iloc[0] if not df_cogs.empty else 0.0,2), str(y2):round(df_cogs[str(y2)].iloc[0] if not df_cogs.empty else 0.0,2), str(y3):round(df_cogs[str(y3)].iloc[0] if not df_cogs.empty else 0.0,2)},
        {"rubrica":"FSE", str(y1):round(df_fse[str(y1)].iloc[0] if not df_fse.empty else 0.0,2), str(y2):round(df_fse[str(y2)].iloc[0] if not df_fse.empty else 0.0,2), str(y3):round(df_fse[str(y3)].iloc[0] if not df_fse.empty else 0.0,2)},
        {"rubrica":"Pessoal", str(y1):round(tot_pess[str(y1)],2), str(y2):round(tot_pess[str(y2)],2), str(y3):round(tot_pess[str(y3)],2)},
        {"rubrica":"Depreciações", str(y1):round(dep_tot[str(y1)],2), str(y2):round(dep_tot[str(y2)],2), str(y3):round(dep_tot[str(y3)],2)},
        {"rubrica":"Juros", str(y1):round(juros_tot[str(y1)],2), str(y2):round(juros_tot[str(y2)],2), str(y3):round(juros_tot[str(y3)],2)},
    ])
    df_dr.loc[len(df_dr)] = {"rubrica":"Resultado",
                             str(y1): round(df_dr[str(y1)].iloc[0]-df_dr[str(y1)].iloc[1]-df_dr[str(y1)].iloc[2]-df_dr[str(y1)].iloc[3]-df_dr[str(y1)].iloc[4]-df_dr[str(y1)].iloc[5],2),
                             str(y2): round(df_dr[str(y2)].iloc[0]-df_dr[str(y2)].iloc[1]-df_dr[str(y2)].iloc[2]-df_dr[str(y2)].iloc[3]-df_dr[str(y2)].iloc[4]-df_dr[str(y2)].iloc[5],2),
                             str(y3): round(df_dr[str(y3)].iloc[0]-df_dr[str(y3)].iloc[1]-df_dr[str(y3)].iloc[2]-df_dr[str(y3)].iloc[3]-df_dr[str(y3)].iloc[4]-df_dr[str(y3)].iloc[5],2)}

    inv_total = sum([float(it.get("valor",0)) for it in dados.get("investimento",[])])
    bal = pd.DataFrame([
        {"rubrica":"Ativo Não Corrente", str(y1): round(inv_total,2), str(y2): None, str(y3): None},
        {"rubrica":"Ativo Corrente", str(y1): round(max(0.1*tot[str(y1)],0.0),2), str(y2): round(max(0.1*tot[str(y2)],0.0),2), str(y3): round(max(0.1*tot[str(y3)],0.0),2)},
        {"rubrica":"Capital Próprio", str(y1): round(float(dados.get("capitais_proprios_iniciais",0.0)),2), str(y2): None, str(y3): None},
    ])

    return {"vendas":df_vendas,"cogs":df_cogs,"fse":df_fse,"pessoal":df_pessoal,"depreciacoes":df_dep,"financiamento":df_fin,"dr":df_dr,"balanco":bal}

def build_docx(cfg: Dict[str, Any], tabs: Dict[str, pd.DataFrame], out_path: Path):
    doc = Document()
    doc.add_heading("Formulário IEFP — Preenchimento Automático (via Site)", 0)

    idt = cfg.get("identificacao", {})
    doc.add_heading("Identificação", 1)
    for k in ["designacao_social","nif","promotor","forma_juridica","morada","email","telefone","cae"]:
        p = doc.add_paragraph()
        p.add_run(k.replace("_"," ").title()+": ").bold=True
        p.add_run(str(idt.get(k,"")))

    limites = cfg.get("limites_caracteres", {"objetivos_projeto":2000,"mercado":1200,"instalacoes":1000})
    tx = cfg.get("textos", {})
    from autofill_core import aplicar_limite as _cut
    doc.add_heading("Objetivos do Projeto", 1)
    doc.add_paragraph(_cut(tx.get("objetivos_projeto",""), limites.get("objetivos_projeto",2000)))
    doc.add_heading("Mercado", 1)
    doc.add_paragraph(_cut(tx.get("mercado",""), limites.get("mercado",1200)))
    doc.add_heading("Instalações", 1)
    doc.add_paragraph(_cut(tx.get("instalacoes",""), limites.get("instalacoes",1000)))

    def df_to_table(title, df):
        doc.add_heading(title, 2)
        if df is None or df.empty:
            doc.add_paragraph("(sem dados)"); return
        rows, cols = df.shape
        table = doc.add_table(rows=rows+1, cols=cols); table.style='Table Grid'
        for j, col in enumerate(df.columns): table.cell(0,j).text = str(col)
        for i in range(rows):
            for j in range(cols):
                val = df.iat[i,j]
                table.cell(i+1, j).text = "" if pd.isna(val) else str(val)

    df_to_table("Vendas/Serviços (3 anos)", tabs["vendas"])
    df_to_table("COGS (3 anos)", tabs["cogs"])
    df_to_table("FSE (3 anos)", tabs["fse"])
    df_to_table("Pessoal (3 anos)", tabs["pessoal"])
    df_to_table("Depreciações", tabs["depreciacoes"])
    df_to_table("Financiamento", tabs["financiamento"])
    df_to_table("Demonstração de Resultados", tabs["dr"])
    df_to_table("Balanço", tabs["balanco"])

    doc.save(out_path)
